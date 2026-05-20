"""Build the Member 3 presentation script DOCX.

ChatLog Link: https://chatgpt.com/share/69fac10c-8ebc-839c-a19d-c83718097045
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs/slides/member3_presentation_script_qa.md"
OUT = ROOT / "docs/slides/member3_presentation_script_qa.docx"

IMPORTANT_PATTERNS = [
    "Balanced Geo + NN + 2-opt",
    "47.661 percent",
    "47.661%",
    "47.66 percent",
    "47.66%",
    "61.80 percent",
    "61.80%",
    "56.24 percent",
    "56.24%",
    "472.537 kilometers",
    "472.537 km",
    "247.324 kilometers",
    "247.324 km",
    "1692.977 minutes",
    "871.263 minutes",
    "exactly 10 orders per rider",
    "not a full CVRP solver",
    "common depot",
    "Haversine distance",
    "traffic and weather",
    "capacity feasibility",
    "prescriptive changes",
]


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(88, 49, 0)

    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(17)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(88, 49, 0)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    h2 = styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(88, 49, 0)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    h3 = styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(67, 67, 67)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(3)

    return doc


def add_text_with_emphasis(paragraph, text: str) -> None:
    spans: list[tuple[int, int]] = []
    for pattern in IMPORTANT_PATTERNS:
        for match in re.finditer(re.escape(pattern), text, flags=re.IGNORECASE):
            spans.append((match.start(), match.end()))
    for match in re.finditer(r"`([^`]+)`|\*\*([^*]+)\*\*", text):
        spans.append((match.start(), match.end()))
    spans.sort()

    merged: list[tuple[int, int]] = []
    for start, end in spans:
        if not merged or start > merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    pos = 0
    for start, end in merged:
        if start > pos:
            paragraph.add_run(text[pos:start])
        token = text[start:end]
        token = token.strip("`")
        if token.startswith("**") and token.endswith("**"):
            token = token[2:-2]
        run = paragraph.add_run(token)
        run.bold = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text_with_emphasis(p, text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    doc.add_paragraph()


def build() -> None:
    doc = setup_doc()
    title = doc.add_paragraph(style="Title")
    title.add_run("Member 3 Presentation Script and Q&A Prep")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Target speaking time: about 3 minutes. Key numbers and answer anchors are bolded.").italic = True

    md = SRC.read_text(encoding="utf-8")
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_emphasis(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_text_with_emphasis(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("If asked for"):
            doc.add_paragraph(line, style="Heading 3")
        elif line.startswith("Member 3 implemented"):
            add_callout(doc, line)
        else:
            p = doc.add_paragraph()
            add_text_with_emphasis(p, line)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
